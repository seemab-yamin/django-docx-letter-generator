import base64
import json
import os

from django.conf import settings
from django.http import FileResponse, HttpResponseNotFound
from django.shortcuts import render
from docx import Document


def json_to_docx(data, docx_file):
    """
    Converts a JSON file into a formatted DOCX file.

    Args:
        data (dict): Json Data
        docx_file (str): Path to the output DOCX file.
    """

    doc = Document()

    # Add a title
    doc.add_heading("JSON Data in DOCX", 0)

    # Iterate through the data and add it to the document
    for key, value in data.items():
        p = doc.add_paragraph()
        p.add_run(f"{key}: ").bold = True
        p.add_run(str(value))

    doc.save(docx_file)


def download_file(request, encoded_path):
    # Decode the Base64-encoded path
    file_path = base64.urlsafe_b64decode(encoded_path.encode()).decode()

    if os.path.exists(file_path):
        with open(file_path, "rb") as f:
            f = open(file_path, "rb")
            response = FileResponse(
                f,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{os.path.basename(file_path)}"',
            )
            return response
    else:
        return HttpResponseNotFound(f"File not found:{file_path}")


def render_form(request):
    if request.method == "POST":
        submitted_data = request.POST.dict()
        print("Submitted Data:", submitted_data)
        file_name = "form_results.docx"
        file_path = os.path.join(settings.MEDIA_ROOT, file_name)

        submitted_data.pop("csrfmiddlewaretoken")
        json_to_docx(submitted_data, file_path)

        return render(request, "download.html", context={"file_name": file_name})
    else:
        input_file_path = os.path.join(
            settings.BASE_DIR, "form_handler", settings.INPUT_FORM_TEMPLATE
        )
        with open(input_file_path, "r") as file:
            form_templates = json.load(file)

        return render(
            request, "render_form.html", context={"form_templates": form_templates}
        )
